import streamlit as st
import openai
import os
from pydub import AudioSegment
from tempfile import NamedTemporaryFile
import math

OPENAI_API_KEY = st.secrets["api_key"]
import streamlit as st
from openai import OpenAI
from docx import Document
import json
import os

# Set up OpenAI API key
key = st.secrets["api_key"]
client = OpenAI(api_key=key)

# Prompts for the OpenAI API
system_prompt = "Translate the following text to English. Provide only the translation with no comments or notes."

# Function to call OpenAI's translation model with chunked text
def translate_text(text_chunk):
    chat_completion = client.chat.completions.create(
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text_chunk}
        ],
        model="gpt-4o",  # Use the correct model
        temperature=0.2,
        max_tokens=3000,
        top_p=0.1,
        frequency_penalty=0.2,
        presence_penalty=0.1,
        stop=None
    )

    response_json = json.loads(chat_completion.model_dump_json(indent=2))
    content = response_json['choices'][0]['message']['content']
    return content.strip()  # Strip any extra whitespace to ensure clean output

# Function to read the .docx file and extract text
def read_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to split text into chunks (approximately 1500 tokens per chunk)
def split_text_into_chunks(text, max_chunk_size=1500):
    words = text.split()
    chunks = []
    chunk = []
    chunk_size = 0

    for word in words:
        word_length = len(word) + 1  # Account for the space between words
        if chunk_size + word_length > max_chunk_size:
            chunks.append(' '.join(chunk))
            chunk = [word]
            chunk_size = word_length
        else:
            chunk.append(word)
            chunk_size += word_length

    # Add the last remaining chunk
    if chunk:
        chunks.append(' '.join(chunk))

    return chunks

# Function to save the translated text into a .docx file
def save_translated_text_to_docx(translated_text, output_path):
    doc = Document()
    doc.add_paragraph(translated_text)
    doc.save(output_path)

# Function to generate the new filename with "_translation" appended
def get_translated_filename(uploaded_file):
    file_name = os.path.splitext(uploaded_file.name)[0]  # Get the file name without extension
    return f"{file_name}_translation.docx"  # Append "_translation" and add the .docx extension

# Streamlit App
def main():
    st.title("Document Translation to English")

    # File uploader for .docx files
    uploaded_file = st.file_uploader("Upload a .docx file", type="docx")

    if uploaded_file is not None:
        # Step 1: Read the uploaded .docx file
        original_text = read_docx(uploaded_file)

        # Step 2: Split the text into manageable chunks
        chunks = split_text_into_chunks(original_text)

        # Step 3: Display the number of chunks
        num_chunks = len(chunks)
        st.write(f"The document will be translated in **{num_chunks} chunks**.")

        # Step 4: Translate when the user clicks the button
        if st.button("Translate"):
            translated_texts = []
            for i, chunk in enumerate(chunks):
                # Translate each chunk
                translated_chunk = translate_text(chunk)
                translated_texts.append(translated_chunk)

                # Display each chunk as it's translated
                st.write(f"Chunk {i + 1} out of {num_chunks}:")
                st.text_area(f"Translated Chunk {i + 1}", translated_chunk, height=150)

            # Combine all the translated chunks into a single text
            final_translated_text = '\n'.join(translated_texts)

            # Step 5: Generate translated file name
            output_file_name = get_translated_filename(uploaded_file)

            # Save the translated text to a new .docx file
            save_translated_text_to_docx(final_translated_text, output_file_name)

            # Display complete translated text
            st.text_area("Complete Translated Text", final_translated_text, height=400)

            # Step 6: Offer the option to download the translated document
            with open(output_file_name, "rb") as f:
                st.download_button(
                    label="Download Translated Document",
                    data=f,
                    file_name=output_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
